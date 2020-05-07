# -*- coding: utf-8 -*-
"""
Created on Thu Apr 30 21:48:14 2020

@author: vidis
"""

"""
------
Inputs
------

1) Locations, Latitudes & Longitudes
2) Number of days of stay 

-------
Outputs
-------

{'Day1':[Locations in chronological order]}
{'Day2':[Locations in chronological order]}
{'Day3':[Locations in chronological order]}


"""

"""
Heuristic Definitions for the A*Search 
--------------------------------------

1) F(x) = f(x) + g(x)
    
f(x) = Least distance score for specific combination & specific permutation
g(x) = Total distance to all the other locations pending to visit from the hotel (straight line distance so that the heuristic is admissible)

---------------------------------------

Start state = [0*len(locations)]
End State = [1*len(locations)]

"""

"""
Function 5 --> get the valid comb and assess each of the valid combinatiosn

Input --> (0,2,3)

Output --> total shortest distance (based on permutations and hotel)
"""

    
"""
IMPLEMENTATION OF A* SEARCH

Assuming that there are 9 locations, start state will be a matrix of 9 0's

%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%
START NODE = [0,0,0,0,0,0...]              %%
State 0 means that the node is not visited %%
%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%

%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%
GOAL NODE = [1,1,1,1,1,1,1....]                %%
State 1 means that all nodes have been visited %%
%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%%

GOAL TEST
--> To check if all states have been visited using the A* Search


OPEN LIST 

CLOSED LIST


===========
HEURISTICS
===========

H(X) = F(X) + G(X)

F(X) --> Least distance obtained using Genetic Algorithm 
G(X) --> Unexplored states distance to center of origin

"""

from math import radians, sqrt, cos, sin, asin
from pymongo import MongoClient
import pandas as pd
import numpy as np
from itertools import combinations
from itertools import permutations
import math
from pymongo import MongoClient
import dns
from docx import Document
import os
import ipyleaflet
from ipyleaflet import Map, Marker, MarkerCluster
from ipyleaflet import Map, basemaps, basemap_to_tiles, Rectangle, Circle
import pandas as pd
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders
import smtplib
from .StaticMap import *
from docx.enum.text import WD_BREAK
from .GeneticAlgoSearchSpace import geneticAlgorithm


def statesNotExplored(index, comb):
    index2 = []
    for element in index:
        if element in comb:
            pass
        else:
            index2.append(element)
    
    return index2

def DistanceNotExplored(index, Hotel_list,Attractions):
    distanceNotExplored = 0
    for item in index:
        locNotExplored = [(Attractions.iloc[item,1],Attractions.iloc[item,2])]
        distanceNotExplored = distanceNotExplored + Distance(Hotel_list[0][1],Hotel_list[0][0],locNotExplored[0][1],locNotExplored[0][0])
    
    return distanceNotExplored


def Distance(lon1, lat1, lon2, lat2):
    """    
    Distance is calculated between two points using radians traversed along the 
    radius of the earth. 
    
    Radius of Earth in km is 6371 km
    """
    lon1, lat1, lon2, lat2 = map(radians, [lon1, lat1, lon2, lat2])
    
    dlon = lon2 - lon1 
    dlat = lat2 - lat1 
    
    a = sin(dlat/2)**2 + cos(lat1) * cos(lat2) * sin(dlon/2)**2
    c = 2 * asin(sqrt(a)) 
    
    distance = 6371* c
    return distance


def currentAttractions(current_state, Attractions):
    comb = [i for i, e in enumerate(current_state) if e == 1]
    for item in comb:
        Attractions = Attractions.drop(item)
    return Attractions



def Comb(current_state, Attractions, combinations_to_analyse):
    Comb = []
    for length in combinations_to_analyse:
        comb = combinations(Attractions.index.tolist(), length)
        print(comb)
        for i in comb:
            print(i)
            Comb.append(i)
            
    return Comb



def CumulativeStates(start_state, objects):
    position = objects.position
    for item in position:
        start_state[item]=1
    return start_state


#cum_list = [[0,0,1],[0,0,0],[0,1,0]]
def cumulativeStates(start_state, cum_list):
    for items in cum_list:
        index = [i for i, e in enumerate(items) if e ==1]
        print(index)
        for item in index:  
            start_state[item] = 1
    
    return start_state

#cum_state = cumulativeStates(start_state, cum_list)
# a = cumulativeStates([1,0,1],[0,1,0])


def cumulativestates(start_state,cum_state):
    index = [i for i , e in enumerate(cum_state) if e ==1]
    for item in index:
        start_state[item]=1
    return start_state

# a = cumulativestates([1,0,1],[0,1,0])


def convert_to_tuple(list1):
    list2 = []
    list3 = []
    for i in list1:
        list2.append(i)
        t = tuple(list2)
        list3.append(t)
        
    return list3
  


def tuple_to_int(test_tuple):
    return int(''.join(map(str, test_tuple))) 



def validComb(Comb,Attractions):
    valid_comb = []
    for comb in Comb:
        totalTime = 0
        for element in comb:
            time = Attractions.iloc[element,3]
            totalTime = time + totalTime
        if totalTime > 13:
            pass
        else:
            valid_comb.append(comb)
    
            
    return valid_comb


def comb_to_state(start_state, comb):
    for element in comb:
        if isinstance(element, str):
            element = int(element)
        start_state[element] = 1
    return start_state

def state_to_comb(state):
    comb = [i for i, e in enumerate(state) if e == 1]
    comb = tuple(comb)
    return comb



def checkCurrentState(parent,child):
    for element in child:
        if element == 1:
            element1 = child.index
            parent[element1]=1
    return parent



class Node_A():
    """A node class for A* Pathfinding"""

    def __init__(self, parent=None, position=None):
        self.parent = parent
        self.position = position

        self.g = 0
        self.h = 0
        self.f = 0

    def __eq__(self, other):
        return self.position == other.position

ScoreList = []
def Astar(Attractions, Hotel_list, start, end):
    """Returns a list of tuples as a path from the given start to the given end in the given maze"""

    # Create start and end node
    """
    Find the start_node.h and start_node.f/
    End node f, g and h values are defined as 0 until the path is found in which case it is updated
    
    """
    start_node = Node(None, start)
    start_node.h = DistanceNotExplored(Attractions.index.tolist(),Hotel_list,Attractions)
    start_node.g = 0
    start_node.f = start_node.g + start_node.h
        
    end_node = Node(None, end)
    end_node.g = end_node.h = end_node.f = 0

    """
    Initialize both open and closed list 
    
    OpenList --> Fringe List which contains all the nodes that are being explored and associated with a parent node
    ClosedList --> Expanded nodes in the open list are saved in the closed list 
    
    ClosedList --> [(parent_nodes)(child_node)]

    """

    open_list = []
    closed_list = []

    """
    OpenList save first node (parent node 1 --> start node)
    """
    open_list.append(start_node)
    ScoreList = []

    # Loop until you find the end
    while len(open_list) > 0:
        
        """
        contains all the parent nodes and the child nodes
        """
        # Get the current node
        current_node = open_list[0]
        current_index = 0
        for index, item in enumerate(open_list):
            if item.f < current_node.f:
                current_node = item
                current_index = index
                
                        
        """
        Check which node to pop from the open list and append to the closed list.
        Loop through all the nodes in the open list and choose the node with the lowest f score.
        """

        # Pop current off open list, add to closed list
        open_list.pop(current_index)
        closed_list.append(current_node)

        # Found the goal
        if current_node == end_node:
            path = []
            current = current_node
            while current is not None:
                path.append(current.position)
                current = current.parent
            return path[::-1] # Return reversed path
        


        # Generate children
        children = []
        
        Attractions_filtered = currentAttractions(current_node.position, Attractions)
        if len(Attractions_filtered) == 1:
            valid_child_combs = convert_to_tuple(Attractions_filtered.index.tolist())
        else:
            child_combs = Comb(current_node.position, Attractions_filtered, combinations_to_analyse=[2])
            valid_child_combs = validComb(child_combs,Attractions)
        for child_comb in valid_child_combs:
            path_list = []
            for element in child_comb:
                Latitude = Attractions.iloc[element,1]
                Longitude = Attractions.iloc[element,2]
                attraction_loc = (float(Latitude), float(Longitude))
                path_list.append(attraction_loc)
            current_state = comb_to_state([0]*len(Attractions), child_comb)
            child_node = Node(current_node.position, current_state)
            children.append(child_node)            
                
        # Loop through children
        for child in children:
            # Child is on the closed list
            for closed_child in closed_list:
                if child == closed_child:
                    continue
            
            child_comb = state_to_comb(child.position)
            path_list = []
            for element in child_comb:
                Latitude = Attractions.iloc[element,1]
                Longitude = Attractions.iloc[element,2]
                attraction_loc = (float(Latitude), float(Longitude))
                path_list.append(attraction_loc)
            popSize = math.ceil(0.5*len(path_list))
            eliteSize = math.ceil(0.5*popSize)
            mutationRate = 0.01
            generations = 300
            bestRoute, child.g = geneticAlgorithm(path_list, popSize, eliteSize, mutationRate, generations, Hotel_list)
            index = statesNotExplored(Attractions.index.tolist(),child_comb)
            child.h = DistanceNotExplored(index,Hotel_list,Attractions)
            child.f = child.g + child.h
            CombScore = [child_comb, bestRoute, child.position, child.g, child.h, child.f]
            ScoreList.append(CombScore)

            # Child is already in the open list
            for open_node in open_list:
                if child == open_node and child.g > open_node.g:
                    continue

            # Add the child to the open list
            open_list.append(child)



"""
Greedy Best First Search 

computes only based on f(x) = h(x)

"""

    
class Node():
    """A node class for A* Pathfinding"""

    def __init__(self, parent=None, position=None):
        self.parent = parent
        self.position = position
        self.h = 0


def GreedyBestFirstSearch(Attractions, Hotel_list, start, end, duration, combinations_to_analyse):
    """Returns a list of tuples as a path from the given start to the given end in the given maze"""

    # Create start and end node
    """
    Find the start_node.h and start_node.f/
    End node f, g and h values are defined as 0 until the path is found in which case it is updated
    
    """
    start_node = Node(None, start)
    start_node.h = 0
        
    end_node = Node(None, end)
    end_node.h = 0

    """
    Initialize both open and closed list 
    
    OpenList --> Fringe List which contains all the nodes that are being explored and associated with a parent node
    ClosedList --> Expanded nodes in the open list are saved in the closed list 
    
    ClosedList --> [(parent_nodes)(child_node)]

    """

    open_list = []
    closed_list = []


    """
    OpenList save first node (parent node 1 --> start node)
    """
    open_list.append(start_node)
    PathFinder = []
    ScoreList = []
    Attractions_filtered = Attractions

    # Loop until you find the end
    while len(open_list) > 0:
        
        day = len(closed_list)+1
        
        """
        contains all the parent nodes and the child nodes
        """
        # Get the current node
        current_node = open_list[0]
        current_index = 0

        print(closed_list)
        for index, item in enumerate(open_list):
            if item.h <= current_node.h:
                current_node = item
                print(current_node.position)
                current_index = index
                PathFinder.append(current_node.position)
                break
                
        open_list = []
    
        """
        Check which node to pop from the open list and append to the closed list.
        Loop through all the nodes in the open list and choose the node with the lowest f score.
        """

        # Reinitialise the openlist
        closed_list.append(current_node)
        

        if len(closed_list)==1:
            cumulative_states=[0]*len(Attractions)
        else:
            cumulative_states = [0]*len(Attractions)
            cum_list = []
            for item in closed_list:
                cum_list.append(item.position)
            for item in cum_list:
                index = [i for i,e in enumerate(item) if e==1]
                for i in index:
                    cumulative_states[i]=1
                    
        print("CumulativeStates: %s on day %s"%(cumulative_states,day))

        # Found the goal
        if cumulative_states == end_node.position:
            return PathFinder, ScoreList      
        
        
        print("Start Day %s"%(day))
        # Generate children
        children = []
        
        Attractions_filtered = currentAttractions(current_node.position, Attractions_filtered)
        print(Attractions_filtered)
        print("Its an issue here")
        print(day)
        print(duration)
        if day==duration:
            print("Its going here")
            if len(Attractions_filtered.index.tolist()) == 1:
                valid_child_combs = convert_to_tuple(Attractions_filtered.index.tolist())
                valid_child_combs = tuple(valid_child_combs[0])
                valid_child_combs = [str(x) for x in valid_child_combs]
            else:
                valid_child_combs = [tuple(Attractions_filtered.index.tolist())]
        else:
            print("Its going here lalalalal")
            print(current_node.position)
            print(combinations_to_analyse)
            child_combs = Comb(current_node.position, Attractions_filtered, combinations_to_analyse)
            valid_child_combs = validComb(child_combs,Attractions)
            print(valid_child_combs)
            
        i = 0    
        for child_comb in valid_child_combs:
            path_list = []
            i = i +1
            print(i/len(valid_child_combs))
            if len(Attractions_filtered)==1:
                element = int(child_comb)
                Latitude = Attractions.iloc[element,1]
                Longitude = Attractions.iloc[element,2]
                attraction_loc = (float(Latitude), float(Longitude))
                path_list.append(attraction_loc)
                current_state = comb_to_state([0]*len(Attractions), valid_child_combs)
                child_node = Node(current_node.position, current_state)
                children.append(child_node)
            else:
                for element in child_comb:
                    Latitude = Attractions.iloc[element,1]
                    Longitude = Attractions.iloc[element,2]
                    attraction_loc = (float(Latitude), float(Longitude))
                    path_list.append(attraction_loc)
                current_state = comb_to_state([0]*len(Attractions), child_comb)
                child_node = Node(current_node.position, current_state)
                children.append(child_node)   

        i =0
        # Loop through children
        for child in children:
            i=i+1
            print(i/len(valid_child_combs))
            # Child is on the closed list
            for closed_child in closed_list:
                if child == closed_child:
                    continue
            print("Running child nodes on day%s"%(day))
            child_comb = state_to_comb(child.position)
            path_list = []
            for element in child_comb:
                Latitude = Attractions.iloc[element,1]
                Longitude = Attractions.iloc[element,2]
                attraction_loc = (float(Latitude), float(Longitude))
                path_list.append(attraction_loc)
            popSize = math.ceil(0.5*len(path_list))
            eliteSize = math.ceil(0.5*popSize)
            mutationRate = 0.01
            generations = 300
            bestRoute, child.h = geneticAlgorithm(path_list, popSize, eliteSize, mutationRate, generations, Hotel_list)
            CombScore = [child_comb, bestRoute, child.position, child.h]
            ScoreList.append(CombScore)

            # Child is already in the open list
            for open_node in open_list:
                if child == open_node and child.h > open_node.h:
                    continue

            # Add the child to the open list
            open_list.append(child)


def removeExtraSquareBrackets(list1):
    list2 = []
    for items in list1[0]:
        list2.append(items)
    return list2


def AttractionsToVisit(username):
    
    Mongourl = "mongodb+srv://vidish:tripatus@cluster0-jzyrn.mongodb.net/test"
    
    client = MongoClient(Mongourl)
    
    d = dict((db, [collection for collection in client[db].collection_names()])
        for db in client.database_names())
    
    db = client.TravelPlan
    
    collection = db.PlacesToVisit
    
    Attractions = pd.DataFrame()
    
    lat = []
    lng = []
    loc = []
    
    for doc in collection.find({'username':username}):
        loc.append(doc['Locations'])
        lat.append(doc['Latitude'])
        lng.append(doc['Longitude'])
        
    Attractions['Location'] = flatten(loc)
    Attractions['Latitude'] = flatten(lat)
    Attractions['Longitude'] = flatten(lng)
    
    Hotel = pd.DataFrame()
    
    Hotel['Hotel'] = doc['Hotel']
    Hotel['Latitude'] = doc['Hotel_Latitude']
    Hotel['Longitude'] = doc['Hotel_Longitude']

    Hotel_list = [(float(doc['Hotel_Latitude'][0]), float(doc['Hotel_Longitude'][0]))]
    
    return Attractions, doc['Duration of trip'], Hotel, Hotel_list

# Attractions, duration, Hotel, Hotel_list = AttractionsToVisit('Tour')

def CombinationsToAnalyse(Attractions, Duration_of_visit):
    Loc_perday = math.ceil(len(Attractions)/Duration_of_visit)
    combinations_to_analyse = [Loc_perday]
    if 0 in combinations_to_analyse:
        combinations_to_analyse.remove(0)
    elif 1 in combinations_to_analyse:
        combinations_to_analyse.remove(1)
    return combinations_to_analyse

# combinations_to_analyse = CombinationsToAnalyse(Attractions, duration)

def timeSpent(Attractions):
    
    Mongourl = "mongodb+srv://vidish:tripatus@cluster0-jzyrn.mongodb.net/test"
    
    client = MongoClient(Mongourl)
    
    d = dict((db, [collection for collection in client[db].collection_names()])
        for db in client.database_names())

    db = client.TravelPlan
    
    collection = db.AttractionDB
    
    
    print("-----------------CONNECTION SUCCESSFUL----------------------")
    
    Attractions_list = []
    Categories = []
    Likeability_Solo = []
    Likeability_Family = []
    Likeability_Friends = []
    Time_Spent = []
    Latitude = []
    Longitude = []
    Description = []
    Image_Path = []
        
    for doc in collection.find():
        Attractions_list.append(doc['PAGETITLE'])
        Categories.append(doc['Leaf Node Category'])
        Likeability_Solo.append(doc['Likeability Solo'])
        Likeability_Family.append(doc['Likeability Family'])
        Likeability_Friends.append(doc['Likeability Friends'])
        Time_Spent.append(doc['Hour'])
        Latitude.append(doc['Latitude'])
        Longitude.append(doc['Longitude'])
        Description.append(doc['Description'])
        Image_Path.append(doc['Image Path'])
    
    print(Attractions)
    
    Likeability_Df = pd.DataFrame()
    
    Likeability_Df['Attractions'] = Attractions_list
    Likeability_Df['Attraction Category'] = Categories
    Likeability_Df['Popularity (solo)'] = [float(x) for x in Likeability_Solo]
    Likeability_Df['Popularity (Family)'] = [float(x) for x in Likeability_Family]
    Likeability_Df['Popularity (Friends)'] = [float(x) for x in Likeability_Friends]
    Likeability_Df['Time Spent'] = [float(x) for x in Time_Spent]
    Likeability_Df['Latitude'] = [float(x) for x in Latitude]
    Likeability_Df['Longitude'] = [float(x) for x in Longitude]
    Likeability_Df['Description'] = Description
    Likeability_Df['Image Path'] = Image_Path

    print(Likeability_Df)
    TimeSpent = []
    for locations in Attractions['Location'].tolist():
        temp_df = Likeability_Df[Likeability_Df['Attractions'] == locations]
        TimeSpent.append(temp_df['Time Spent'].tolist())
    print(TimeSpent)
    print(len(TimeSpent))
    Attractions['Time Spent'] = flatten(TimeSpent)
    
    return Attractions



def Path_Evaluator(path, score):
    score_df = pd.DataFrame.from_records(score)
    Distance_score = []
    path_taken = []
    for index, each_day_path in enumerate(path):
        if index == 0:
            continue
        else:
            print(index)
            score_list = []
            path_list = []
            for index, rows in score_df.iterrows():
                if rows[2] == each_day_path:
                    score_list.append(rows[3])
                    path_list.append(rows[1])
            for x,y in zip(score_list,path_list):
                if x == min(score_list):
                    print(x)
                    Distance_score.append(x)
                    path_taken.append(y)
                break
    return Distance_score, path_taken


def DailyPathsUpdate(username, Attractions, path_taken):
    DayPaths = []
    for daily_path in path_taken:
        Daily_path_locs = []
        for cds in daily_path:
            for index, rows in Attractions.iterrows():
                if (rows['Latitude'] == cds[0]) and (rows['Longitude'] == cds[1]):
                    Daily_path_locs.append(rows['Location'])
        DayPaths.append(Daily_path_locs)
    travelPlan = []
    for index, eachDayPath in enumerate(DayPaths):
        travelPlan.append({"Day"+str(index+1): eachDayPath})
    TravelPlan = {'travelPlan':travelPlan}
    Mongourl = "mongodb+srv://vidish:tripatus@cluster0-jzyrn.mongodb.net/test"
    
    client = MongoClient(Mongourl)
    
    d = dict((db, [collection for collection in client[db].collection_names()])
        for db in client.database_names())
    
    db = client.TravelPlan
    
    collection = db.PlacesToVisit
    
    collection.update({'username':username},{'$set':TravelPlan},upsert=True)
        
    return DayPaths


def removeExtraBrackets(list1):
    list2 = []
    for item in list1[0]:
        list2.append(item)
    
    return list2



def flatten(list2):
    list3 = []
    for sublist in list2:
        for item in sublist:
            list3.append(item)
            
    return list3


def convert_cds_to_tuple(Latitudes, Longitudes):
    cds_tuple = []
    for x,y in zip(Latitudes, Longitudes):
        cds = [x,y]
        cds = tuple(cds)
        cds_tuple.append(cds)
    
    return cds_tuple


        
def Attractions_Db():

    Mongourl = "mongodb+srv://vidish:tripatus@cluster0-jzyrn.mongodb.net/test"
        
    client = MongoClient(Mongourl)
    
    db = client.TravelPlan
    
    collection = db.AttractionDB
    
    Locations = []
    Latitudes = []
    Longitudes = []
    for item in collection.find():
        Locations.append(item['PAGETITLE'])
        Latitudes.append(item['Latitude'])
        Longitudes.append(item['Longitude'])
        
    Location = pd.DataFrame()
    Location['LocationsList'] = Locations
    Location['Latitudes'] = Latitudes
    Location['Longitudes'] = Longitudes
        
    return Location


def UserTravelData(username):
    
    Mongourl = "mongodb+srv://vidish:tripatus@cluster0-jzyrn.mongodb.net/test"
        
    client = MongoClient(Mongourl)
    
    db = client.TravelPlan
    
    collection = db.PlacesToVisit

    travelPlan = []
    Hotel = []
    Hotel_Latitude = []
    Hotel_Longitude = []
    
    for doc in collection.find({'username': username}):
        travelPlan.append(doc['travelPlan'])
        travelPlan = travelPlan[0]
        Hotel.append(doc['Hotel'])
        Hotel_Latitude.append([float(x) for x in doc['Hotel_Latitude']])
        Hotel_Longitude.append([float(x) for x in doc['Hotel_Longitude']])
    
    return travelPlan, Hotel, flatten(Hotel_Latitude), flatten(Hotel_Longitude)



def squareList(list2):
    list3=[]
    for sublist in list2:
        if sublist is None:
            pass
        else:
            for item in sublist:
                list3.append(item)
    
    return list3



def createDocument(username):

    document = Document()
    
    document.add_heading("Trip@Us - Your Friendly Travel Planner \n")
    
    document.add_paragraph("Hello %s, please find your curated itineary for your Singapore visit below :). \n \n\
    We provide a detailed day by day breakdown of your trip to some of the places we know you will love. And what's more, our core algorithm optimised your travelling distance so that you can spend more time at the places you love :)"%(username))
    
    Attraction_Df = Attractions_Db()
    travelPlan, Hotel, Hotel_Latitude, Hotel_Longitude = UserTravelData(username)
    
    Mongourl = "mongodb+srv://vidish:tripatus@cluster0-jzyrn.mongodb.net/test"
        
    client = MongoClient(Mongourl)
    
    db = client.TravelPlan
    
    collection = db.AttractionsDB
    
    i = 0
    
    while i < len(travelPlan):
        day = i+1
        document.add_heading("Day %s"%(day))
        Latitudes = [Hotel_Latitude]
        Longitudes = [Hotel_Longitude]
        for item in travelPlan[i]['Day'+str(day)]:
            document.add_paragraph(item, style='List Bullet')
            Latitude = Attraction_Df[Attraction_Df['LocationsList']==item]['Latitudes'].tolist()
            Longitude = Attraction_Df[Attraction_Df['LocationsList']==item]['Longitudes'].tolist()
            Latitudes.append([float(x) for x in Latitude])
            Longitudes.append([float(x) for x in Longitude])
            Latitudes.append(Hotel_Latitude)
            Longitudes.append(Hotel_Longitude)
        
        Latitudes = squareList(Latitudes)
        Longitudes = squareList(Longitudes)
        cds_tuple = convert_cds_to_tuple(Longitudes,Latitudes)
        
        map = StaticMap(300, 300, 10)
        line = Line(cds_tuple, 'blue', 3)
        for x,y in zip(Hotel_Longitude, Hotel_Latitude):
            Hotel_marker = tuple([x,y])
            print(Hotel_marker)
        marker_outline = CircleMarker(Hotel_marker, 'white', 18)
        marker = CircleMarker(Hotel_marker, '#0036FF', 12)
        map.add_marker(marker)
        map.add_marker(marker_outline)
        map.add_line(line)
        image = map.render()
        image.save('TourPlan.png')
        
        document.add_picture('TourPlan.png')
                # print(cds_tuple)
    
        if i == len(travelPlan)-1:
            pass
        else:
            document.add_page_break()
    
    
        i = i +1
            
    
    
    print(os.getcwd())
    document.save("TravelPlan_Trip@Us.docx")
    



def send_email(user_email, username):
    mail = smtplib.SMTP("smtp.gmail.com", 587)
    mail.ehlo()
    mail.starttls()
    msg = MIMEMultipart()
    from_msg = 'tripatusnus@gmail.com'
    msg['From'] = 'tripatusnus@gmail.com'
    msg['To'] = user_email
    msg['Subject'] = "Trip@Us Itinerary Report"
    body = "Hi %s, \n \n Please find attached your curated itinerary for your visit to Singapore. \n \n We sincerely hope you enjoy your trip to Singapore and while you are there do check out our dynamic planning telebot travelBot for planning on the go :) \n \n Thank you, \n Trip@Us Team"%(username)
    msg.attach(MIMEText(body,'plain'))
    filename = 'Trip@Us Itinerary Report.pdf'
    attachment=open("TravelPlan_Trip@Us.docx",'rb')
    part = MIMEBase('application', 'octet-stream')
    part.set_payload((attachment).read())
    encoders.encode_base64(part)
    part.add_header('Content-Disposition',"attachment; filename = "+filename)
    msg.attach(part)
    text = msg.as_string()
    mail.login('tripatusnus@gmail.com','TripAtUs2020')
    mail.sendmail(from_msg,user_email,text)
    mail.close()




def plan_my_trip(username, user_email):
    Attractions, duration, Hotel, Hotel_list = AttractionsToVisit(username)
    print(duration)
    combinations_to_analyse = CombinationsToAnalyse(Attractions, duration)
    Attractions = timeSpent(Attractions)
    start_state = [0]*len(Attractions)
    end_state = [1]*len(Attractions)
    path, score = GreedyBestFirstSearch(Attractions, Hotel_list, start_state,end_state,duration,combinations_to_analyse)
    Distance_score, path_taken = Path_Evaluator(path,score)
    DayPaths = DailyPathsUpdate(username, Attractions, path_taken)
    createDocument(username)
    send_email(user_email, username)

# for testing
# plan_my_trip('park','park@gmail.com')



